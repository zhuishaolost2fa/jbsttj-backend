"""Word / Office 文档（.docx）文本提取。

.. versionchanged:: 修复标题检测

   旧版仅靠 Word 内置样式名（Heading 1 / 标题 1）判定标题，且把所有块的
   ``font_size`` 留 0.0 —— 导致两个连锁问题：

   1. 大量 Word 手册（尤其是非技术用户排版的）不使用内置标题样式，
      而是手动加粗 / 放大字号来区分标题。旧版把这类段落全部判成 ``body``，
      ``build_section_paths`` 拿不到任何 heading 块，所有 chunk 与 QA 的
      ``section_path`` 和 ``title`` 全部为空。

   2. ``font_size`` 全为 0 时，``calibrate_headings`` 检测到 ``sizes`` 为空
      直接 return，跳过全局标题层级校准 —— 连字号倍率这条路也堵死了。

   修复策略（三级降级）：

   a. **样式名**（最可靠）—— ``Heading 1`` / ``标题 2`` 等；
   b. **字号 + 文本特征**（主修复）—— 从 ``run.font.size`` 提取字号，
      填入 ``font_size`` 与 ``font_sizes``，让 ``calibrate_headings`` 正常工作；
   c. **文本模式兜底**（补刀）—— 样式与字号都没信号时，用与 PDF 相同的
      ``_HEADING_PATTERNS`` 正则（``第X章`` / ``X.X 标题``）做最后兜底。

   三级合起来覆盖了「有样式」「有字号差异但无样式」「无字号差异但有编号前缀」
   三种 Word 手册排版场景。

.docx 本质是 OOXML（ZIP 包里的 XML），**自带文字层**，无需 OCR。
这里把它解析成与 :mod:`app.services.pdf_extract` 同构的
:class:`~app.services.pdf_extract.TextBlock` / :class:`~app.services.pdf_extract.ShardResult`，
于是下游的 `merge_shards` / `strip_noise` / `calibrate_headings` / `build_section_paths` /
`chunk_blocks` 全部**零改动复用**——Word 手册同样能切出带章节面包屑的语义块、同样能去重、同样能向量化。

**伪页码**
下游一整套逻辑都基于 ``page`` 字段（页眉页脚剥离靠 y_ratio + 出现频次、
跨分片续接靠 ``head.page - tail.page <= 1``、检索结果展示「出处第 N 页」）。
Word 没有真页码，这里给每个块打一个**伪页码** ``global_index // blocks_per_page + 1``：
同一窗口内的相邻块伪页码相同，跨分片边界的相邻块伪页码差 ``<= 1``，
于是 ``merge_shards`` 的跨片续接仍然成立，``total_pages`` 也落在一个合理的量级。

模块内**不做任何 IO 之外的业务判断**，纯函数式，方便离线测试。
"""

from __future__ import annotations

import logging
import math
import re
import statistics
from typing import Any, Dict, Iterator, List, Optional, Tuple

from app.services.pdf_extract import ShardResult, TextBlock, _is_heading_like, _SENTENCE_END

logger = logging.getLogger("app.doc_extract")

# python-docx 是可选依赖：没装时本模块的函数调用会抛清晰的错误，
# 但其余依赖它的 import 不会在模块加载期就崩。
try:  # pragma: no cover - 取决于运行环境
    import docx

    _HAS_DOCX = True
except ImportError:  # pragma: no cover
    docx = None  # type: ignore[assignment]
    _HAS_DOCX = False


# 标题样式名里的层级提取，命中「Heading 1」「标题 2」「Heading 3」等
_HEADING_RE = re.compile(r"(?:heading|标题)\s*(\d+)|(\d+)", re.IGNORECASE)


def _heading_level_from_style(style_name: Optional[str]) -> int:
    """从 Word 样式名推断标题层级（1-6），非标题返回 0。

    样式命名五花八门，按优先级匹配：
      - ``Heading 1`` / ``标题 1`` / ``Heading 3`` -> 对应数字
      - ``Title`` / ``标题`` -> 1
      - ``Subtitle`` / ``副标题`` -> 2
    数字超出 6 的夹紧到 6（极深层级对章节面包屑无意义）。
    """
    if not style_name:
        return 0
    name = style_name.strip()
    low = name.lower()

    m = _HEADING_RE.search(name)
    if m:
        lvl = int(m.group(1) or m.group(2))
        return max(1, min(6, lvl))

    if low in ("title", "标题"):
        return 1
    if low in ("subtitle", "副标题"):
        return 2
    return 0


def _para_is_bold(paragraph: Any) -> bool:
    """段落整体是否加粗（任一 run 显式加粗即算）。"""
    try:
        runs = paragraph.runs
    except Exception:  # noqa: BLE001 - 个别异常段落不阻断整体
        return False
    if not runs:
        return False
    return any(bool(getattr(r, "bold", False)) for r in runs)


def _extract_font_size(paragraph: Any) -> float:
    """从段落的 runs 中提取主导字号（字符加权中位数）。

    python-docx 的 ``run.font.size`` 返回 ``Pt`` 对象（``Length`` 子类），
    显式设置过字号的 run 才有值；继承自样式的 run 返回 ``None``。

    取字符加权中位数而非简单平均：正文段落字多、权重高，
    确保中位数稳定落在正文字号上，与 PDF 侧的统计口径一致。
    """
    sizes: List[float] = []
    try:
        runs = paragraph.runs
    except Exception:  # noqa: BLE001
        return 0.0
    if not runs:
        return 0.0
    for r in runs:
        text = getattr(r, "text", "") or ""
        if not text.strip():
            continue
        size = getattr(getattr(r, "font", None), "size", None)
        if size is not None:
            # size 是 Emu/Length 对象，转 float 得到 Pt 值
            pt = float(size)
            if pt > 0:
                sizes.extend([pt] * len(text))
    if not sizes:
        return 0.0
    return statistics.median(sizes)


def _heading_level_from_text(text: str, *, is_bold: bool = False) -> int:
    """从文本特征推断标题层级（样式与字号都没信号时的兜底）。

    复用 PDF 侧的 ``_is_heading_like`` 正则匹配（第X章 / X.X / 标题N 等），
    叠加长度与加粗约束，避免把正文里碰巧带编号的句子误判成标题。

    返回的层级与 ``classify_block`` 的 level 4-5 区间对齐，
    ``calibrate_headings`` 会做层级压缩，这里给一个合理初始值即可。
    """
    stripped = text.strip()
    if not stripped or len(stripped) > 40:
        return 0
    if not _is_heading_like(stripped):
        return 0
    # 以句末标点结尾的不太可能是标题（标题通常不以句号收尾）
    if _SENTENCE_END.search(stripped) is not None:
        return 0
    # 加粗的编号标题 → level 4；不加粗但短 → level 5
    if is_bold:
        return 4
    if len(stripped) <= 24:
        return 5
    return 0


def open_document(path: str):
    """打开 .docx，返回 python-docx Document。"""
    if not _HAS_DOCX:
        raise RuntimeError(
            "未安装 python-docx，无法解析 Word 文档。请执行：pip install python-docx"
        )
    return docx.Document(path)


def probe_unit_count(path: str) -> int:
    """统计可提取的文本单元总数（段落 + 表格行），用于规划分片。

    与 :func:`app.services.pdf_extract.probe_page_count` 对应：PDF 返回页数，
    这里返回「单元数」——Word 的分片单位不是页而是文本单元。
    """
    doc = open_document(path)
    try:
        count = len(doc.paragraphs)
        for table in doc.tables:
            count += len(table.rows)
        return count
    finally:
        # python-docx 在 Document 构造时就把 zip 解包到内存，不持有文件句柄，
        # 无 close 方法；del 让它尽快被 GC 即可。
        del doc


def _iter_units(doc: Any) -> Iterator[Tuple[str, Any]]:
    """把文档扁平化成 (kind, obj) 序列：段落优先，随后每张表的每一行。

    ``kind`` 为 ``"p"``（段落）或 ``"t"``（表格行），``obj`` 是对应的
    python-docx 对象。调用方负责维护全局序号（global index），
    分片规划与跨片续接都依赖这个连续序号。
    """
    for p in doc.paragraphs:
        yield "p", p
    for table in doc.tables:
        for row in table.rows:
            yield "t", row


def extract_shard(
    path: str,
    *,
    shard_index: int,
    unit_start: int,
    unit_end: int,
    blocks_per_page: int = 40,
) -> ShardResult:
    """提取 [unit_start, unit_end] 闭区间（1-based，全局单元序号）的文本块。

    与 :func:`app.services.pdf_extract.extract_shard` 的契约一致：
    返回 ``ShardResult``，下游 ``merge_shards`` 等无需感知来源格式。

    Word 文档不做 OCR（本来就有文字层）。表格行被并成一条 ``table`` 块，
    单元格之间用 `` | `` 分隔，便于后续向量化时保留表格结构。
    """
    doc = open_document(path)
    try:
        blocks: List[TextBlock] = []
        font_sizes: List[float] = []
        page_start = 0
        page_end = 0

        for gi, (kind, obj) in enumerate(_iter_units(doc), start=1):
            if gi < unit_start:
                continue
            if gi > unit_end:
                break

            if kind == "p":
                style_name = obj.style.name if getattr(obj, "style", None) else ""
                level = _heading_level_from_style(style_name)
                text = (obj.text or "").strip()
                is_bold = _para_is_bold(obj)
                font_size = _extract_font_size(obj)

                # 样式没识别为标题时，用文本特征兜底：
                # 先看字号（让 calibrate_headings 全局校准），
                # 字号也没信号时用正则匹配编号前缀
                if level == 0 and text:
                    level = _heading_level_from_text(text, is_bold=is_bold)

                is_heading = level > 0
            else:  # 表格行
                cells = [c.text.strip() for c in obj.cells]
                text = " | ".join(c for c in cells if c)
                is_heading = False
                level = 0
                is_bold = False
                font_size = 0.0

            if not text:
                continue

            page = (gi - 1) // max(1, blocks_per_page) + 1
            if page_start == 0:
                page_start = page
            page_end = page

            # 字号样本按字符数加权收集，供 calibrate_headings 全局统计
            if font_size > 0:
                weight = min(20, max(1, len(text) // 10))
                font_sizes.extend([font_size] * weight)

            blocks.append(
                TextBlock(
                    text=text,
                    page=page,
                    block_type="heading" if is_heading else "body",
                    heading_level=level if is_heading else 0,
                    font_size=font_size,
                    is_bold=bool(is_bold),
                    y_ratio=0.5,  # 无版面坐标，落在页面中部，页眉页脚剥离不会误伤
                )
            )

        page_count = max(0, page_end - page_start + 1)
        return ShardResult(
            shard_index=shard_index,
            page_start=page_start,
            page_end=page_end,
            blocks=blocks,
            font_sizes=font_sizes,
            page_count=page_count,
        )
    finally:
        del doc
